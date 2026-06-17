from pathlib import Path
from datetime import datetime
import re

import pdfplumber
import pytesseract
from docx import Document

from agents.translation_agent import TranslationAgent


"""
    Extractor Agent
    Persona: Document Parser
    Uses LLM: No
    Flow:
        extractor_agent.py -> translation_agent.py
"""

class ExtractorAgent:
   

    #Keeps a list of supported file types.
    #Creates a TranslationAgent instance.
    def __init__(self):
        self.supported_extensions = [".pdf", ".docx", ".png"]
        self.translation_agent = TranslationAgent()

   


    #it main entry point mentod triggeres from monitor_agent
    # Input: file path + optional email metadata.
    #Check file exists → if not, return failure.
    #Check extension → must be .pdf, .docx, or .png.
    #Extract raw text using helper methods:
    def extract(self, file_path: str, email_metadata: dict | None = None) -> dict:
        email_metadata = email_metadata or {}
        path = Path(file_path)

        if not path.exists():
            return {
                "status": "failed",
                "file_path": file_path,
                "email_metadata": email_metadata,
                "error": "File not found"
            }

        file_extension = path.suffix.lower()

        if file_extension not in self.supported_extensions:
            return {
                "status": "failed",
                "file_path": file_path,
                "email_metadata": email_metadata,
                "error": f"Unsupported file type: {file_extension}"
            }

        if file_extension == ".pdf":
            raw_text = self._extract_pdf(path)
        elif file_extension == ".docx":
            raw_text = self._extract_docx(path)
        elif file_extension == ".png":
            raw_text = self._extract_png(path)
        else:
            raw_text = ""

        invoice_data = self._extract_structured_data(raw_text)

        state = {
            "status": "extracted",
            "file_path": str(path),
            "file_name": path.name,
            "file_type": file_extension.replace(".", ""),
            "email_metadata": email_metadata,
            "raw_text": raw_text,
            "invoice_data": invoice_data,
            "audit_trail": [
                {
                    "agent": "Extractor Agent",
                    "persona": "Document Parser",
                    "status": "success",
                    "timestamp": datetime.utcnow().isoformat(),
                    "message": "Document text and structured fields extracted."
                }
            ]
        }

        return self.translation_agent.translate(state)
    
    #Uses pdfplumber to read each page of file.
    #Collects text blocks → joins them into one string.
    def _extract_pdf(self, path: Path) -> str:
        text_blocks = []

        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text_blocks.append(page.extract_text() or "")

        return "\n".join(text_blocks).strip()

    #Uses python-docx -Reads paragraphs + tables.
    #Tables are flattened into | separated values.
    def _extract_docx(self, path: Path) -> str:
        document = Document(str(path))
        text_blocks = []

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_blocks.append(paragraph.text.strip())

        for table in document.tables:
            for row in table.rows:
                row_values = []

                for cell in row.cells:
                    if cell.text.strip():
                        row_values.append(cell.text.strip())

                if row_values:
                    text_blocks.append(" | ".join(row_values))

        return "\n".join(text_blocks).strip()
    

    #Uses pytesseract (OCR), Converts image text into string.
    def _extract_png(self, path: Path) -> str:
        return pytesseract.image_to_string(str(path)).strip()

    """Extracts key invoice fields: -Invoice No,Invoice Date (normalized to YYYY-MM-DD),
    Vendor ID / Vendor Name,PO Number,Currency (USD, EUR, INR, GBP),
    Total Amount,Line Items (SKU, description, 
    qty, unit price, total)"""
    def _extract_structured_data(self, raw_text: str) -> dict:
        invoice_date_raw = self._find_value(
            raw_text,
            [
                r"Invoice\s*Date\s*[:\-]\s*([0-9]{4}-[0-9]{2}-[0-9]{2})",
                r"Invoice\s*Date\s*[:\-]\s*([0-9]{2}/[0-9]{2}/[0-9]{4})",
                r"Invoice\s*Date\s*[:\-]\s*([0-9]{1,2}\s+[A-Za-z]+\s+[0-9]{4})",
                r"Date\s*[:\-]\s*([0-9]{1,2}\s+[A-Za-z]+\s+[0-9]{4})"
            ]
        )

        line_items = self._extract_line_items(raw_text)

        return {
            "invoice_no": self._find_value(
                raw_text,
                [
                    r"Invoice\s*(?:No|Number|#)\s*[:\-]\s*([A-Za-z0-9\-]+)",
                    r"invoice_no\s*[:\-]\s*([A-Za-z0-9\-]+)"
                ]
            ),
            "invoice_date": self._normalize_date(invoice_date_raw),
            "vendor_id": self._find_value(
                raw_text,
                [
                    r"Vendor\s*ID\s*[:\-]\s*([A-Za-z0-9\-]+)",
                    r"vendor_id\s*[:\-]\s*([A-Za-z0-9\-]+)"
                ]
            ),
            "vendor_name": self._find_value(
                raw_text,
                [
                    r"Vendor\s*Name\s*[:\-]\s*(.+)",
                    r"Vendor\s*[:\-]\s*(.+)"
                ]
            ),
            "po_number": self._find_value(
                raw_text,
                [
                    r"PO\s*(?:No|Number|#)\s*[:\-]\s*([A-Za-z0-9\-]+)",
                    r"PO\s*Reference\s*[:\-]\s*([A-Za-z0-9\-]+)",
                    r"Purchase\s*Order\s*[:\-]\s*([A-Za-z0-9\-]+)"
                ]
            ),
            "currency": self._find_currency(raw_text),
            "total_amount": self._find_total_amount(raw_text, line_items),
            "line_items": line_items
        }


    #converts multiple formats to standard
    def _normalize_date(self, date_text: str | None) -> str | None:
        if not date_text:
            return None

        date_text = date_text.strip()

        date_formats = [
            "%Y-%m-%d",
            "%d/%m/%Y",
            "%d %B %Y",
            "%d %b %Y"
        ]

        for date_format in date_formats:
            try:
                parsed_date = datetime.strptime(date_text, date_format)
                return parsed_date.strftime("%Y-%m-%d")
            except ValueError:
                continue

        return date_text

    #regex match for single field
    def _find_value(self, text: str, patterns: list):
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)

            if match:
                return match.group(1).strip()

        return None

    #detect currency codes/symbols
    def _find_currency(self, text: str):
        match = re.search(r"\b(USD|EUR|INR|GBP)\b", text, re.IGNORECASE)

        if match:
            return match.group(1).upper()

        if "$" in text:
            return "USD"
        if "€" in text:
            return "EUR"
        if "₹" in text:
            return "INR"
        if "£" in text:
            return "GBP"

        return None

    #regex or sum of line items.
    def _find_total_amount(self, text: str, line_items: list):
        patterns = [
            r"Total\s*Amount\s*[:\-]\s*([0-9]+(?:\.[0-9]+)?)",
            r"Grand\s*Total\s*[:\-]\s*([0-9]+(?:\.[0-9]+)?)",
            r"Invoice\s*Total\s*[:\-]\s*([0-9]+(?:\.[0-9]+)?)"
        ]

        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)

            if match:
                return float(match.group(1))

        if line_items:
            return round(
                sum(float(item.get("total") or 0) for item in line_items),
                2
            )

        return None

    #parse SKU-based rows
    def _extract_line_items(self, text: str) -> list:
        line_items = []

        for line in text.splitlines():
            clean_line = line.strip()

            if not clean_line:
                continue

            item_match = re.search(r"\bSKU-[A-Za-z0-9]+\b", clean_line)

            if not item_match:
                continue

            item_code = item_match.group(0)

            if "|" in clean_line:
                parts = [part.strip() for part in clean_line.split("|")]

                if len(parts) >= 5:
                    try:
                        line_items.append(
                            {
                                "item_code": parts[0],
                                "description": parts[1],
                                "qty": float(parts[2]),
                                "unit_price": float(parts[3]),
                                "total": float(parts[4])
                            }
                        )
                        continue
                    except ValueError:
                        pass

            qty = self._find_number_after_label(
                clean_line,
                ["Qty", "Quantity"]
            )

            unit_price = self._find_number_after_label(
                clean_line,
                ["Unit Price", "UnitPrice", "Price", "Unit"]
            )

            total = self._find_number_after_label(
                clean_line,
                ["Total", "Line Total"]
            )

            description = clean_line.replace(item_code, "").strip()

            line_items.append(
                {
                    "item_code": item_code,
                    "description": description,
                    "qty": qty,
                    "unit_price": unit_price,
                    "total": total
                }
            )

        return line_items


    #regex to capture numbers after labels like "Qty" or "Unit Price".
    def _find_number_after_label(self, text: str, labels: list):
        for label in labels:
            pattern = rf"{label}\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)"
            match = re.search(pattern, text, re.IGNORECASE)

            if match:
                return float(match.group(1))

        return None